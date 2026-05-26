#!/usr/bin/env python3
"""Generate a realistic sample 2B ukeplan .docx for testing the scraper.

Run:  python samples/make_sample_ukeplan.py
Then: python scraper.py --file samples/ukeplan-eksempel.docx
"""
from pathlib import Path

from docx import Document

OUT = Path(__file__).with_name("ukeplan-eksempel.docx")

doc = Document()
doc.add_heading("Vekeplan 2B - veke 22", level=1)
doc.add_paragraph("Rosseland skule")
doc.add_paragraph("Ei kjekk veke! Hugs gymtøy på tysdag og symjetøy på torsdag.")

timeplan = doc.add_table(rows=5, cols=4)
rows = [
    ["Dag", "Økt 1", "Økt 2", "Økt 3"],
    ["Måndag", "Norsk", "Matematikk", "Bibliotek"],
    ["Tysdag", "Gym", "Norsk", "Kunst og handverk"],
    ["Onsdag", "Matematikk", "Engelsk", "Norsk"],
    ["Torsdag", "Symjing", "Matematikk", "Musikk"],
]
for r, cells in enumerate(rows):
    for c, val in enumerate(cells):
        timeplan.rows[r].cells[c].text = val

doc.add_paragraph("")
doc.add_paragraph("Lekser denne veka:")
lekser = doc.add_table(rows=4, cols=2)
lekse_rows = [
    ["Til tysdag", "Norsk: les side 24-25. Matte: oppgåve 5 og 6."],
    ["Til onsdag", "Engelsk: øv på glosene (colours)."],
    ["Til torsdag", "Lesebok: les 10 minutt høgt heime."],
    ["Til fredag", "Ingen lekse - god helg!"],
]
for r, cells in enumerate(lekse_rows):
    for c, val in enumerate(cells):
        lekser.rows[r].cells[c].text = val

doc.add_paragraph("")
doc.add_paragraph("Måndag: ta med matpakke til turdag.")

doc.save(OUT)
print(f"Wrote {OUT}")
