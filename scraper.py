#!/usr/bin/env python3
"""Fetch the weekly plan (ukeplan) for class 2B at Rosseland skole and email it.

The plan is published as a Word .docx linked from the class section page. The
file name embeds the ISO week number but ends in a per-week random hash, so the
URL cannot be predicted - we load the section page, pick the link for the
current week, download that .docx and parse it.

Designed to run on GitHub Actions (which has internet access to minskole.no).
"""
from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import smtplib
import ssl
import sys
import time
from dataclasses import dataclass, asdict
from datetime import datetime
from email.message import EmailMessage
from io import BytesIO
from pathlib import Path
from urllib.parse import urljoin
from zoneinfo import ZoneInfo

import requests
from bs4 import BeautifulSoup
from docx import Document

SECTION_URL = "https://www.minskole.no/rosseland/seksjon/23538"
CLASS_NAME = "2B"
SCHOOL_NAME = "Rosseland"
TZ = ZoneInfo("Europe/Oslo")
USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
)
STATE_PATH = Path("data/state.json")
DEBUG_DIR = Path("debug")

# Words worth flagging in the weekly summary (Norwegian).
HIGHLIGHTS = {
    "gym": "Gym/PE",
    "kroppsøving": "Gym/PE",
    "symjing": "Swimming",
    "svømming": "Swimming",
    "symje": "Swimming",
    "bibliotek": "Library",
    "lekse": "Homework",
    "prøve": "Test",
    "tur": "Trip",
    "matpakke": "Packed lunch",
}

WEEK_RE = re.compile(r"veke[\s\-_]*(\d{1,2})", re.IGNORECASE)


@dataclass
class PlanLink:
    week: int | None
    url: str
    text: str


def make_session() -> requests.Session:
    s = requests.Session()
    s.headers.update({"User-Agent": USER_AGENT, "Accept-Language": "nb,no,en;q=0.8"})
    return s


def get(session: requests.Session, url: str, *, retries: int = 4) -> requests.Response:
    """GET with exponential backoff on network errors."""
    delay = 2
    last_exc: Exception | None = None
    for attempt in range(retries):
        try:
            resp = session.get(url, timeout=30)
            resp.raise_for_status()
            return resp
        except requests.RequestException as exc:
            last_exc = exc
            if attempt < retries - 1:
                time.sleep(delay)
                delay *= 2
    raise RuntimeError(f"Failed to GET {url}: {last_exc}")


def find_plan_links(html: str, base_url: str) -> list[PlanLink]:
    """Find all 'Vekeplan' .docx links on the section page."""
    soup = BeautifulSoup(html, "lxml")
    links: list[PlanLink] = []
    for a in soup.find_all("a", href=True):
        href = a["href"]
        text = a.get_text(strip=True)
        haystack = f"{href} {text}".lower()
        if "vekeplan" not in haystack and "ukeplan" not in haystack:
            continue
        if ".docx" not in href.lower() and ".doc" not in href.lower():
            # Some links may point to a document without an obvious extension.
            if "document" not in href.lower():
                continue
        m = WEEK_RE.search(href) or WEEK_RE.search(text)
        week = int(m.group(1)) if m else None
        links.append(PlanLink(week=week, url=urljoin(base_url, href), text=text))
    return links


def select_plan(links: list[PlanLink], target_week: int) -> PlanLink | None:
    if not links:
        return None
    for link in links:
        if link.week == target_week:
            return link
    weeks = [l for l in links if l.week is not None]
    if weeks:
        return max(weeks, key=lambda l: l.week)
    return links[0]


def parse_docx(data: bytes) -> dict:
    """Extract paragraphs and tables from the .docx into a structured dict."""
    doc = Document(BytesIO(data))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    return {"paragraphs": paragraphs, "tables": tables}


def plan_to_text(plan: dict) -> str:
    """Render parsed plan content into readable plain text."""
    lines: list[str] = []
    for para in plan["paragraphs"]:
        lines.append(para)
    for table in plan["tables"]:
        lines.append("")
        for row in table:
            cells = [c for c in row if c]
            if cells:
                lines.append("  " + " | ".join(cells))
    return "\n".join(lines).strip()


def find_highlights(text: str) -> list[str]:
    found: dict[str, None] = {}
    low = text.lower()
    for word, label in HIGHLIGHTS.items():
        if word in low:
            found[label] = None
    return list(found.keys())


def content_hash(text: str) -> str:
    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def render_email(week: int, plan_text: str, plan_url: str) -> tuple[str, str]:
    subject = f"Ukeplan {CLASS_NAME} - veke {week} ({SCHOOL_NAME})"
    highlights = find_highlights(plan_text)
    parts = [f"Ukeplan for {CLASS_NAME} - veke {week}", ""]
    if highlights:
        parts.append("Hugs denne veka: " + ", ".join(highlights))
        parts.append("")
    parts.append(plan_text or "(Fann ikkje tekstinnhald i dokumentet.)")
    parts.append("")
    parts.append(f"Kjelde: {plan_url}")
    return subject, "\n".join(parts)


def send_email(subject: str, body: str) -> None:
    user = os.environ.get("GMAIL_USER")
    password = os.environ.get("GMAIL_APP_PASSWORD")
    recipient = os.environ.get("RECIPIENT", "thomasmoi87@gmail.com")
    if not user or not password:
        raise RuntimeError(
            "GMAIL_USER and GMAIL_APP_PASSWORD must be set to send email."
        )
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = user
    msg["To"] = recipient
    msg.set_content(body)
    context = ssl.create_default_context()
    with smtplib.SMTP("smtp.gmail.com", 587, timeout=30) as server:
        server.starttls(context=context)
        server.login(user, password)
        server.send_message(msg)
    print(f"Sent email to {recipient}: {subject}")


def load_state() -> dict:
    if STATE_PATH.exists():
        return json.loads(STATE_PATH.read_text(encoding="utf-8"))
    return {}


def save_state(state: dict) -> None:
    STATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    STATE_PATH.write_text(json.dumps(state, indent=2, ensure_ascii=False), encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description="Rosseland 2B weekly plan scraper")
    parser.add_argument("--dump", action="store_true",
                        help="Discovery mode: save section HTML + .docx to debug/, no email")
    parser.add_argument("--force", action="store_true",
                        help="Send email even if the plan is unchanged")
    parser.add_argument("--no-email", action="store_true",
                        help="Parse and print, but do not send email or update state")
    parser.add_argument("--week", type=int, default=None,
                        help="Override the target ISO week number")
    args = parser.parse_args()

    target_week = args.week or datetime.now(TZ).isocalendar().week
    session = make_session()

    print(f"Fetching section page: {SECTION_URL}")
    page = get(session, SECTION_URL)
    links = find_plan_links(page.text, SECTION_URL)
    print(f"Found {len(links)} candidate plan link(s): "
          + ", ".join(f"veke {l.week}" for l in links) if links else "Found 0 plan links")

    if args.dump:
        DEBUG_DIR.mkdir(parents=True, exist_ok=True)
        (DEBUG_DIR / "section.html").write_text(page.text, encoding="utf-8")
        print(f"Saved section HTML to {DEBUG_DIR / 'section.html'}")

    if not links:
        title = BeautifulSoup(page.text, "lxml").title
        print("No 'Vekeplan' links found. The page may require login or have changed.")
        print(f"Page title: {title.get_text(strip=True) if title else '(none)'}")
        return 2

    chosen = select_plan(links, target_week)
    print(f"Selected plan: veke {chosen.week} -> {chosen.url}")

    docx_bytes = get(session, chosen.url).content
    if args.dump:
        out = DEBUG_DIR / f"vekeplan-veke-{chosen.week}.docx"
        out.write_bytes(docx_bytes)
        print(f"Saved docx to {out}")

    plan = parse_docx(docx_bytes)
    plan_text = plan_to_text(plan)
    print("\n----- PARSED PLAN -----\n" + plan_text + "\n-----------------------\n")

    if args.dump or args.no_email:
        return 0

    new_hash = content_hash(plan_text)
    state = load_state()
    if state.get("hash") == new_hash and not args.force:
        print("Plan unchanged since last send; skipping email.")
        return 0

    week = chosen.week if chosen.week is not None else target_week
    subject, body = render_email(week, plan_text, chosen.url)
    send_email(subject, body)

    save_state({
        "week": week,
        "hash": new_hash,
        "plan_url": chosen.url,
        "last_sent": datetime.now(TZ).isoformat(),
    })
    print(f"Updated state at {STATE_PATH}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
